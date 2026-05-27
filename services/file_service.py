from io import BytesIO
from pypdf import PdfReader
from docx import Document as DocxDocument


def extract_text_from_file(file_name: str, content_bytes: bytes):
    lower_name = file_name.lower()

    if lower_name.endswith(".txt"):
        return content_bytes.decode("utf-8", errors="ignore")

    if lower_name.endswith(".pdf"):
        reader = PdfReader(BytesIO(content_bytes))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    if lower_name.endswith(".docx"):
        doc = DocxDocument(BytesIO(content_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)

    return None